"""Generate detailed Word document with ALL credentials and 30 cases."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import psycopg2

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Shading Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            tbl.rows[r_idx + 1].cells[c_idx].text = str(val)
    return tbl

# ═══════════════════════════ TITLE ═══════════════════════════
title = doc.add_heading('Prymera — Credenciales del Sistema', level=0)
title.alignment = 1

p = doc.add_paragraph()
p.add_run('Documento completo con todas las credenciales, endpoints y los 30 casos de crédito.\n').italic = True
run = p.add_run('Generado: 02/07/2026')
run.bold = True
run.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════ 1. ARQUITECTURA ═══════════════════════════
doc.add_heading('1. Arquitectura del Sistema', level=1)
doc.add_paragraph(
    'Prymera es un ecosistema de microfinanzas compuesto por:\n\n'
    '• App Fuerza de Ventas (Flutter) — app móvil para asesores de crédito\n'
    '• App Cliente (Flutter) — app móvil para clientes\n'
    '• Frontend Web (Flutter Web) — panel de supervisión\n'
    '• Backend API (FastAPI Python) — puerto 8003\n'
    '• Base de Datos PostgreSQL — puerto 5432, bd_core_mobile'
)

doc.add_paragraph('')

# ═══════════════════════════ 2. POSTGRESQL ═══════════════════════════
doc.add_heading('2. Base de Datos PostgreSQL', level=1)
add_table(doc,
    ['Propiedad', 'Valor'],
    [
        ('Servidor', 'localhost'),
        ('Puerto', '5432'),
        ('Base de datos', 'bd_core_mobile'),
        ('Usuario', 'postgres'),
        ('Contraseña', 'postgres'),
        ('Servicio Windows', 'postgresql-x64-18'),
    ]
)

doc.add_paragraph('')
p = doc.add_paragraph('Para conectar desde pgAdmin: ')
p.add_run('Servers → Register → Server...').bold = True
doc.add_paragraph('Name: bd_core_mobile | Host: localhost | Port: 5432 | Username: postgres | Password: postgres | ✅ Save password')

doc.add_page_break()

# ═══════════════════════════ 3. BACKEND ═══════════════════════════
doc.add_heading('3. Backend API REST (FastAPI)', level=1)
add_table(doc,
    ['Propiedad', 'Valor'],
    [
        ('URL Base', 'http://localhost:8003'),
        ('Swagger UI', 'http://localhost:8003/docs'),
        ('Health Check', 'http://localhost:8003/health'),
        ('Iniciar servidor', 'D:\\backend\\iniciar_api.bat (doble click)'),
        ('Detener servidor', 'Ctrl + C en terminal'),
    ]
)

doc.add_paragraph('')
doc.add_heading('Endpoints principales', level=2)
add_table(doc,
    ['Endpoint', 'Método', 'Descripción'],
    [
        ('/auth/login', 'POST', 'Login de asesores (Fuerza de Ventas)'),
        ('/auth/me', 'GET', 'Perfil del asesor autenticado'),
        ('/cartera/dashboard', 'GET', 'Dashboard del asesor'),
        ('/cartera/diaria', 'GET', 'Cartera diaria asignada'),
        ('/clientes/{id}/ficha', 'GET', 'Ficha completa del cliente'),
        ('/solicitudes', 'POST', 'Registrar nueva solicitud'),
        ('/solicitudes/{id}', 'GET', 'Detalle de solicitud'),
        ('/preevaluacion', 'POST', 'Ejecutar pre-evaluación'),
        ('/buro/consulta', 'POST', 'Consulta a buró'),
        ('/buro/listas', 'POST', 'Consulta listas inhabilitados'),
        ('/reportes/generar', 'GET', 'Reportes gerenciales'),
        ('/cliente_app/login', 'POST', 'Login de clientes'),
        ('/cliente_app/solicitudes', 'POST', 'Crear solicitud desde App Cliente'),
    ]
)

doc.add_page_break()

# ═══════════════════════════ 4. ASESORES ═══════════════════════════
doc.add_heading('4. Fuerza de Ventas — Asesores', level=1)
doc.add_paragraph('Contraseña común para todos los asesores: Prymera2024')
doc.add_paragraph('')

conn = psycopg2.connect(host='localhost', port=5432, dbname='bd_core_mobile', user='postgres', password='postgres')
cur = conn.cursor()
cur.execute("SELECT codigo_empleado, nombres, apellidos, perfil FROM asesores ORDER BY codigo_empleado")
asesores = cur.fetchall()
cur.close()

add_table(doc,
    ['Código', 'Nombre', 'Apellidos', 'Perfil', 'Contraseña'],
    [(a[0], a[1], a[2], a[3], 'Prymera2024') for a in asesores]
)

doc.add_paragraph('')
doc.add_paragraph('La app Fuerza de Ventas se conecta a http://10.0.2.2:8003 (emulador Android).')
doc.add_paragraph('Para cambiar la URL: editar D:\\app_Fventas_Prymera\\lib\\core\\network\\api_client.dart')

doc.add_page_break()

# ═══════════════════════════ 5. APP CLIENTE ═══════════════════════════
doc.add_heading('5. App Cliente — Usuarios', level=1)
doc.add_paragraph('Todos los usuarios tienen contraseña: Prymera2024')
doc.add_paragraph('Username = número de DNI del cliente')
doc.add_paragraph('')

conn = psycopg2.connect(host='localhost', port=5432, dbname='bd_core_mobile', user='postgres', password='postgres')
cur = conn.cursor()
cur.execute("""
    SELECT u.username, c.nombres, c.apellidos, c.numero_documento, c.nombre_negocio
    FROM usuarios_cliente u
    JOIN clientes c ON c.id = u.cliente_id
    ORDER BY u.username
""")
usuarios = cur.fetchall()
cur.close()
conn.close()

add_table(doc,
    ['Usuario (DNI)', 'Nombres', 'Apellidos', 'DNI', 'Negocio', 'Contraseña'],
    [(u[0], u[1], u[2], u[3], u[4] or '—', 'Prymera2024') for u in usuarios]
)

doc.add_page_break()

# ═══════════════════════════ 6. FRONTEND WEB ═══════════════════════════
doc.add_heading('6. Frontend Web — Panel de Supervisión', level=1)
add_table(doc,
    ['Propiedad', 'Valor'],
    [
        ('Proyecto', 'D:\\frontend'),
        ('Tecnología', 'Flutter Web'),
        ('Puerto (dev)', 'http://localhost:51589 (cambia cada vez)'),
        ('Iniciar', 'cd D:\\frontend && flutter run -d chrome'),
        ('Login', 'Simulado — no conecta al backend aún'),
    ]
)

doc.add_paragraph('')
doc.add_paragraph('El login del frontend web es simulado (cualquier usuario/contraseña ingresa).')
doc.add_paragraph('Los supervisores deberían usar las credenciales del asesor con perfil supervisor o administrador:')
add_table(doc,
    ['Código', 'Nombre', 'Perfil', 'Contraseña'],
    [
        ('0003', 'Rosa Condori Apaza', 'supervisor', 'Prymera2024'),
        ('9999', 'Admin Prymera', 'administrador', 'Prymera2024'),
    ]
)

doc.add_page_break()

# ═══════════════════════════ 7. LOS 30 CASOS ═══════════════════════════
doc.add_heading('6. Los 30 Casos de Crédito — Detalle Completo', level=1)

casos = [
    {
        "num": 1, "nombre": "Anaximandro Quispe", "doc": "40118120", "tel": "964110201",
        "negocio": "Bodega Don Anaxi", "tipo": "Comercio", "ubicacion": "El Tambo",
        "ant": 48, "ingreso": 2200, "gasto": 900, "lat": -12.0581, "lng": -75.2027,
        "monto": 1000, "plazo": 12, "tea": "43.92% (sin seguro)", "garantia": "sin garantia",
        "destino": "Capital de trabajo: compra de mercaderia",
        "cuota_ref": 100.95, "prioridad": "normal",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 1 entidad(es), deuda S/4,500, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 1000,
        "fecha_desembolso": "02/02/2026", "dia_pago": 3,
        "cuota_final": 100.95,
    },
    {
        "num": 2, "nombre": "Eulalia Mamani", "doc": "41223341", "tel": "964110202",
        "negocio": "Picanteria La Eulalia", "tipo": "Restaurante", "ubicacion": "Chilca",
        "ant": 36, "ingreso": 3000, "gasto": 1400, "lat": -12.0921, "lng": -75.2105,
        "monto": 3000, "plazo": 12, "tea": "40.92% (con seguro)", "garantia": "sin garantia",
        "destino": "Compra de cocina industrial",
        "cuota_ref": 299.59, "prioridad": "media",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 2 entidad(es), deuda S/12,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 3000,
        "fecha_desembolso": "05/02/2026", "dia_pago": 5,
        "cuota_final": 299.59,
    },
    {
        "num": 3, "nombre": "Teofilo Huaman", "doc": "42330336", "tel": "964110203",
        "negocio": "Maderas Huaman", "tipo": "Carpinteria", "ubicacion": "Pilcomayo",
        "ant": 60, "ingreso": 4200, "gasto": 1800, "lat": -12.0496, "lng": -75.2486,
        "monto": 5000, "plazo": 18, "tea": "43.92% (sin seguro)", "garantia": "sin garantia",
        "destino": "Maquinaria: sierra y cepillo",
        "cuota_ref": 366.02, "prioridad": "media",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 1 entidad(es), deuda S/6,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 5000,
        "fecha_desembolso": "10/02/2026", "dia_pago": 10,
        "cuota_final": 366.02,
    },
    {
        "num": 4, "nombre": "Casandra Flores", "doc": "43440349", "tel": "964110204",
        "negocio": "Distribuidora Casandra", "tipo": "Abarrotes", "ubicacion": "Huancayo",
        "ant": 84, "ingreso": 7000, "gasto": 2600, "lat": -12.0651, "lng": -75.2049,
        "monto": 8000, "plazo": 6, "tea": "43.92% (sin seguro)", "garantia": "sin garantia",
        "destino": "Reposicion de stock por campana",
        "cuota_ref": 1480.73, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 2 entidad(es), deuda S/14,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 8000,
        "fecha_desembolso": "15/02/2026", "dia_pago": 15,
        "cuota_final": 1480.73,
    },
    {
        "num": 5, "nombre": "Demostenes Rojas", "doc": "40556071", "tel": "964110205",
        "negocio": "Ferreteria El Constructor", "tipo": "Ferreteria", "ubicacion": "San Agustin de Cajas",
        "ant": 30, "ingreso": 5200, "gasto": 2100, "lat": -12.0188, "lng": -75.2271,
        "monto": 10000, "plazo": 12, "tea": "43.92% (sin seguro)", "garantia": "hipotecaria",
        "destino": "Ampliacion de local",
        "cuota_ref": 1009.46, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 2 entidad(es), deuda S/12,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 10000,
        "fecha_desembolso": "01/03/2026", "dia_pago": 3,
        "cuota_final": 1009.46,
    },
    {
        "num": 6, "nombre": "Hipatia Condori", "doc": "41669066", "tel": "964110206",
        "negocio": "Confecciones Hipatia", "tipo": "Textil", "ubicacion": "El Tambo",
        "ant": 54, "ingreso": 6800, "gasto": 2900, "lat": -12.0612, "lng": -75.2118,
        "monto": 12000, "plazo": 24, "tea": "40.92% (con seguro)", "garantia": "hipotecaria",
        "destino": "Compra de maquinas remalladoras",
        "cuota_ref": 700.94, "prioridad": "media",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 1 entidad(es), deuda S/6,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 12000,
        "fecha_desembolso": "05/03/2026", "dia_pago": 5,
        "cuota_final": 700.94,
    },
    {
        "num": 7, "nombre": "Anibal Vargas", "doc": "43773379", "tel": "964110207",
        "negocio": "Transportes Anibal", "tipo": "Transporte", "ubicacion": "Concepcion",
        "ant": 42, "ingreso": 9500, "gasto": 4200, "lat": -11.9182, "lng": -75.3142,
        "monto": 15000, "plazo": 18, "tea": "43.92% (sin seguro)", "garantia": "vehicular",
        "destino": "Cuota inicial de vehiculo de carga",
        "cuota_ref": 1098.07, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 2 entidad(es), deuda S/14,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 15000,
        "fecha_desembolso": "10/03/2026", "dia_pago": 10,
        "cuota_final": 1098.07,
    },
    {
        "num": 8, "nombre": "Penelope Apaza", "doc": "40886086", "tel": "964110208",
        "negocio": "Granja Penelope", "tipo": "Avicola", "ubicacion": "Sapallanga",
        "ant": 72, "ingreso": 8800, "gasto": 3600, "lat": -12.1581, "lng": -75.1762,
        "monto": 18000, "plazo": 24, "tea": "43.92% (sin seguro)", "garantia": "hipotecaria",
        "destino": "Ampliacion de galpon",
        "cuota_ref": 1072.10, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 1 entidad(es), deuda S/6,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 18000,
        "fecha_desembolso": "15/03/2026", "dia_pago": 15,
        "cuota_final": 1072.10,
    },
    {
        "num": 9, "nombre": "Heraclito Ccahua", "doc": "41990091", "tel": "964110209",
        "negocio": "Importaciones Heraclito", "tipo": "Comercio", "ubicacion": "Huancayo",
        "ant": 96, "ingreso": 12000, "gasto": 5000, "lat": -12.0668, "lng": -75.2103,
        "monto": 20000, "plazo": 36, "tea": "43.92% (sin seguro)", "garantia": "hipotecaria",
        "destino": "Capital para nueva sucursal",
        "cuota_ref": 927.12, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 2 entidad(es), deuda S/12,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 20000,
        "fecha_desembolso": "02/04/2026", "dia_pago": 3,
        "cuota_final": 927.12,
    },
    {
        "num": 10, "nombre": "Cleopatra Soto", "doc": "43003039", "tel": "964110210",
        "negocio": "Botica Cleopatra", "tipo": "Farmacia", "ubicacion": "Chupaca",
        "ant": 66, "ingreso": 11000, "gasto": 4400, "lat": -12.056, "lng": -75.287,
        "monto": 25000, "plazo": 24, "tea": "40.92% (con seguro)", "garantia": "hipotecaria",
        "destino": "Equipamiento y stock farmaceutico",
        "cuota_ref": 1460.29, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 2 entidad(es), deuda S/14,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 25000,
        "fecha_desembolso": "05/04/2026", "dia_pago": 5,
        "cuota_final": 1460.29,
    },
    {
        "num": 11, "nombre": "Esquilo Ramos", "doc": "40110010", "tel": "964110211",
        "negocio": "Minimarket Esquilo", "tipo": "Bodega", "ubicacion": "Huayucachi",
        "ant": 24, "ingreso": 1900, "gasto": 800, "lat": -12.1339, "lng": -75.2090,
        "monto": 2000, "plazo": 12, "tea": "43.92% (sin seguro)", "garantia": "sin garantia",
        "destino": "Compra de congeladora",
        "cuota_ref": 201.89, "prioridad": "normal",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 1 entidad(es), deuda S/4,500, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 2000,
        "fecha_desembolso": "10/04/2026", "dia_pago": 10,
        "cuota_final": 201.89,
    },
    {
        "num": 12, "nombre": "Ariadna Quispe", "doc": "41226021", "tel": "964110212",
        "negocio": "Estilos Ariadna", "tipo": "Peluqueria", "ubicacion": "El Tambo",
        "ant": 40, "ingreso": 3300, "gasto": 1300, "lat": -12.0573, "lng": -75.2161,
        "monto": 4000, "plazo": 18, "tea": "43.92% (sin seguro)", "garantia": "sin garantia",
        "destino": "Mobiliario y equipos de salon",
        "cuota_ref": 292.82, "prioridad": "media",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 2 entidad(es), deuda S/12,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 4000,
        "fecha_desembolso": "15/04/2026", "dia_pago": 15,
        "cuota_final": 292.82,
    },
    {
        "num": 13, "nombre": "Sofocles Huanca", "doc": "43336033", "tel": "964110213",
        "negocio": "Panaderia Sofocles", "tipo": "Panaderia", "ubicacion": "Sicaya",
        "ant": 58, "ingreso": 5600, "gasto": 2300, "lat": -12.0228, "lng": -75.3134,
        "monto": 6000, "plazo": 12, "tea": "40.92% (con seguro)", "garantia": "sin garantia",
        "destino": "Horno rotativo",
        "cuota_ref": 599.17, "prioridad": "media",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 0 entidad(es), deuda S/0, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 6000,
        "fecha_desembolso": "02/05/2026", "dia_pago": 3,
        "cuota_final": 599.17,
    },
    {
        "num": 14, "nombre": "Casiopea Torres", "doc": "40550055", "tel": "964110214",
        "negocio": "Taller Casiopea", "tipo": "Mecanica", "ubicacion": "Pilcomayo",
        "ant": 50, "ingreso": 7400, "gasto": 3000, "lat": -12.0512, "lng": -75.2451,
        "monto": 7500, "plazo": 6, "tea": "43.92% (sin seguro)", "garantia": "sin garantia",
        "destino": "Herramienta neumatica",
        "cuota_ref": 1388.18, "prioridad": "media",
        "preeval": "APTO (puntaje 85)", "buro": "DEFICIENTE, 2 entidad(es), deuda S/16,000, 45 días mora",
        "decision": "APROBADO", "monto_aprobado": 7500,
        "fecha_desembolso": "05/05/2026", "dia_pago": 5,
        "cuota_final": 1388.18,
    },
    {
        "num": 15, "nombre": "Aristofanes Cruz", "doc": "41669166", "tel": "964110215",
        "negocio": "Insumos Aristofanes", "tipo": "Agropecuario", "ubicacion": "Orcotuna",
        "ant": 78, "ingreso": 8200, "gasto": 3300, "lat": -11.976, "lng": -75.3361,
        "monto": 9000, "plazo": 24, "tea": "43.92% (sin seguro)", "garantia": "hipotecaria",
        "destino": "Capital para campana agricola",
        "cuota_ref": 536.05, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 1 entidad(es), deuda S/6,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 9000,
        "fecha_desembolso": "10/05/2026", "dia_pago": 10,
        "cuota_final": 536.05,
    },
    {
        "num": 16, "nombre": "Calipso Mendoza", "doc": "43880088", "tel": "964110216",
        "negocio": "Calzados Calipso", "tipo": "Calzado", "ubicacion": "Huancayo",
        "ant": 62, "ingreso": 7900, "gasto": 3100, "lat": -12.0689, "lng": -75.2055,
        "monto": 11000, "plazo": 18, "tea": "40.92% (con seguro)", "garantia": "hipotecaria",
        "destino": "Compra de cuero y maquinaria",
        "cuota_ref": 793.03, "prioridad": "media",
        "preeval": "APTO (puntaje 85)", "buro": "CPP, 1 entidad(es), deuda S/9,000, 20 días mora",
        "decision": "APROBADO", "monto_aprobado": 11000,
        "fecha_desembolso": "15/05/2026", "dia_pago": 15,
        "cuota_final": 793.03,
    },
    {
        "num": 17, "nombre": "Demetrio Quispe", "doc": "40119019", "tel": "964110217",
        "negocio": "Mayorista Demetrio", "tipo": "Comercio", "ubicacion": "Jauja",
        "ant": 90, "ingreso": 11500, "gasto": 4700, "lat": -11.7752, "lng": -75.4995,
        "monto": 13500, "plazo": 12, "tea": "43.92% (sin seguro)", "garantia": "hipotecaria",
        "destino": "Reposicion de inventario mayorista",
        "cuota_ref": 1362.77, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 2 entidad(es), deuda S/14,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 13500,
        "fecha_desembolso": "02/06/2026", "dia_pago": 3,
        "cuota_final": 1362.77,
    },
    {
        "num": 18, "nombre": "Antigona Flores", "doc": "41226126", "tel": "964110218",
        "negocio": "Recreo Antigona", "tipo": "Restaurante", "ubicacion": "Concepcion",
        "ant": 70, "ingreso": 9200, "gasto": 3900, "lat": -11.9201, "lng": -75.311,
        "monto": 16000, "plazo": 36, "tea": "43.92% (sin seguro)", "garantia": "hipotecaria",
        "destino": "Ampliacion y remodelacion",
        "cuota_ref": 741.70, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 1 entidad(es), deuda S/6,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 16000,
        "fecha_desembolso": "05/06/2026", "dia_pago": 5,
        "cuota_final": 741.70,
    },
    {
        "num": 19, "nombre": "Pitagoras Rojas", "doc": "43339033", "tel": "964110219",
        "negocio": "Ferreteria Pitagoras", "tipo": "Ferreteria", "ubicacion": "El Tambo",
        "ant": 100, "ingreso": 13000, "gasto": 5200, "lat": -12.0599, "lng": -75.2143,
        "monto": 17000, "plazo": 24, "tea": "40.92% (con seguro)", "garantia": "hipotecaria",
        "destino": "Compra de stock estructural",
        "cuota_ref": 993.00, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 0 entidad(es), deuda S/0, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 17000,
        "fecha_desembolso": "10/06/2026", "dia_pago": 10,
        "cuota_final": 993.00,
    },
    {
        "num": 20, "nombre": "Berenice Apaza", "doc": "40556056", "tel": "964110220",
        "negocio": "Tejidos Berenice", "tipo": "Textil", "ubicacion": "San Jeronimo de Tunan",
        "ant": 46, "ingreso": 8600, "gasto": 3500, "lat": -11.9871, "lng": -75.2899,
        "monto": 19000, "plazo": 18, "tea": "43.92% (sin seguro)", "garantia": "hipotecaria",
        "destino": "Maquinaria de tejido plano",
        "cuota_ref": 1390.89, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 1 entidad(es), deuda S/6,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 19000,
        "fecha_desembolso": "15/06/2026", "dia_pago": 15,
        "cuota_final": 1390.89,
    },
    {
        "num": 21, "nombre": "Anaxagoras Huaman", "doc": "43889089", "tel": "964110221",
        "negocio": "Carga Anaxagoras", "tipo": "Transporte", "ubicacion": "Huancayo",
        "ant": 84, "ingreso": 14000, "gasto": 5800, "lat": -12.0644, "lng": -75.2088,
        "monto": 22000, "plazo": 36, "tea": "43.92% (sin seguro)", "garantia": "vehicular",
        "destino": "Cuota inicial de camion",
        "cuota_ref": 1019.83, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 2 entidad(es), deuda S/14,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 22000,
        "fecha_desembolso": "02/07/2026", "dia_pago": 3,
        "cuota_final": 1019.83,
    },
    {
        "num": 22, "nombre": "Climene Vargas", "doc": "41003001", "tel": "964110222",
        "negocio": "Avicola Climene", "tipo": "Avicola", "ubicacion": "Sapallanga",
        "ant": 76, "ingreso": 13500, "gasto": 5500, "lat": -12.156, "lng": -75.179,
        "monto": 24000, "plazo": 24, "tea": "40.92% (con seguro)", "garantia": "hipotecaria",
        "destino": "Equipamiento de planta",
        "cuota_ref": 1401.88, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 2 entidad(es), deuda S/12,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 24000,
        "fecha_desembolso": "05/07/2026", "dia_pago": 5,
        "cuota_final": 1401.88,
    },
    {
        "num": 23, "nombre": "Epaminondas Soto", "doc": "40115011", "tel": "964110223",
        "negocio": "Bodega Epaminondas", "tipo": "Bodega", "ubicacion": "Pucara",
        "ant": 28, "ingreso": 2600, "gasto": 1000, "lat": -12.1701, "lng": -75.1611,
        "monto": 1500, "plazo": 6, "tea": "43.92% (sin seguro)", "garantia": "sin garantia",
        "destino": "Compra de vitrinas",
        "cuota_ref": 277.64, "prioridad": "normal",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 2 entidad(es), deuda S/12,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 1500,
        "fecha_desembolso": "10/07/2026", "dia_pago": 10,
        "cuota_final": 277.64,
    },
    {
        "num": 24, "nombre": "Lisistrata Ramos", "doc": "41336036", "tel": "964110224",
        "negocio": "Variedades Lisistrata", "tipo": "Comercio", "ubicacion": "Huancayo",
        "ant": 52, "ingreso": 4100, "gasto": 1700, "lat": -12.0633, "lng": -75.2071,
        "monto": 3500, "plazo": 12, "tea": "43.92% (sin seguro)", "garantia": "sin garantia",
        "destino": "Capital de trabajo",
        "cuota_ref": 353.31, "prioridad": "media",
        "preeval": "APTO (puntaje 85)", "buro": "NORMAL, 1 entidad(es), deuda S/6,000, 0 días mora",
        "decision": "APROBADO", "monto_aprobado": 3500,
        "fecha_desembolso": "15/07/2026", "dia_pago": 15,
        "cuota_final": 353.31,
    },
    {
        "num": 25, "nombre": "Filoctetes Cruz", "doc": "41552052", "tel": "964110225",
        "negocio": "Cevicheria Filoctetes", "tipo": "Restaurante", "ubicacion": "Chilca",
        "ant": 18, "ingreso": 3800, "gasto": 2200, "lat": -12.093, "lng": -75.209,
        "monto": 11000, "plazo": 18, "tea": "40.92% (con seguro)", "garantia": "sin garantia",
        "destino": "Ampliacion de local nuevo",
        "cuota_ref": 793.03, "prioridad": "media",
        "preeval": "APTO (puntaje 85)", "buro": "CPP, 2 entidad(es), deuda S/18,000, 15 días mora",
        "decision": "CONDICIONADO", "monto_aprobado": 7000,
        "fecha_desembolso": "02/08/2026", "dia_pago": 3,
        "cuota_final": 504.66,
    },
    {
        "num": 26, "nombre": "Calirroe Mendoza", "doc": "41888088", "tel": "964110226",
        "negocio": "Calzados Calirroe", "tipo": "Calzado", "ubicacion": "El Tambo",
        "ant": 34, "ingreso": 5000, "gasto": 2600, "lat": -12.0588, "lng": -75.2129,
        "monto": 16000, "plazo": 24, "tea": "43.92% (sin seguro)", "garantia": "hipotecaria",
        "destino": "Maquinaria de mayor capacidad",
        "cuota_ref": 952.98, "prioridad": "media",
        "preeval": "APTO (puntaje 85)", "buro": "CPP, 1 entidad(es), deuda S/9,000, 20 días mora",
        "decision": "CONDICIONADO", "monto_aprobado": 10000,
        "fecha_desembolso": "05/08/2026", "dia_pago": 5,
        "cuota_final": 595.61,
    },
    {
        "num": 27, "nombre": "Tucidides Quispe", "doc": "42220022", "tel": "964110227",
        "negocio": "Ferreteria Tucidides", "tipo": "Ferreteria", "ubicacion": "Concepcion",
        "ant": 40, "ingreso": 6200, "gasto": 2900, "lat": -11.9176, "lng": -75.3155,
        "monto": 20000, "plazo": 24, "tea": "40.92% (con seguro)", "garantia": "hipotecaria",
        "destino": "Compra de stock y montacarga",
        "cuota_ref": 1168.23, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "CPP, 2 entidad(es), deuda S/18,000, 15 días mora",
        "decision": "CONDICIONADO", "monto_aprobado": 14000,
        "fecha_desembolso": "10/08/2026", "dia_pago": 10,
        "cuota_final": 817.76,
    },
    {
        "num": 28, "nombre": "Aquiles Mamani", "doc": "43337037", "tel": "964110228",
        "negocio": "Comercial Aquiles", "tipo": "Comercio", "ubicacion": "Huancayo",
        "ant": 60, "ingreso": 9000, "gasto": 3600, "lat": -12.0657, "lng": -75.2099,
        "monto": 15000, "plazo": 24, "tea": "43.92% (sin seguro)", "garantia": "hipotecaria",
        "destino": "Capital de trabajo",
        "cuota_ref": 893.42, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "PERDIDA, 4 entidad(es), deuda S/40,000, 210 días mora, LISTA INHABILITADOS",
        "decision": "RECHAZADO", "monto_aprobado": 0,
        "fecha_desembolso": "—", "dia_pago": "—",
        "cuota_final": "—",
    },
    {
        "num": 29, "nombre": "Medea Apaza", "doc": "41884084", "tel": "964110229",
        "negocio": "Bodega Medea", "tipo": "Bodega", "ubicacion": "Pilcomayo",
        "ant": 22, "ingreso": 1800, "gasto": 1100, "lat": -12.0489, "lng": -75.247,
        "monto": 14000, "plazo": 18, "tea": "43.92% (sin seguro)", "garantia": "sin garantia",
        "destino": "Compra de camioneta para reparto",
        "cuota_ref": 1024.87, "prioridad": "media",
        "preeval": "REVISAR (puntaje 60)", "buro": "DUDOSO, 3 entidad(es), deuda S/25,000, 95 días mora",
        "decision": "RECHAZADO", "monto_aprobado": 0,
        "fecha_desembolso": "—", "dia_pago": "—",
        "cuota_final": "—",
    },
    {
        "num": 30, "nombre": "Esquines Rojas", "doc": "43334034", "tel": "964110230",
        "negocio": "Fletes Esquines", "tipo": "Transporte", "ubicacion": "Jauja",
        "ant": 30, "ingreso": 7000, "gasto": 3200, "lat": -11.774, "lng": -75.501,
        "monto": 30000, "plazo": 24, "tea": "43.92% (sin seguro)", "garantia": "vehicular",
        "destino": "Compra de unidad de transporte",
        "cuota_ref": 1786.83, "prioridad": "alta",
        "preeval": "APTO (puntaje 85)", "buro": "DUDOSO, 3 entidad(es), deuda S/25,000, 95 días mora",
        "decision": "RECHAZADO", "monto_aprobado": 0,
        "fecha_desembolso": "—", "dia_pago": "—",
        "cuota_final": "—",
    },
]

for c in casos:
    doc.add_heading(f'Caso {c["num"]:2d} — {c["nombre"]}', level=2)
    add_table(doc,
        ['Campo', 'Valor'],
        [
            ('Documento (DNI)', c['doc']),
            ('Teléfono', c['tel']),
            ('Negocio', c['negocio']),
            ('Tipo', c['tipo']),
            ('Ubicación', c['ubicacion']),
            ('Antigüedad', f'{c["ant"]} meses'),
            ('Ingreso mensual', f'S/ {c["ingreso"]:,.2f}'),
            ('Gasto mensual', f'S/ {c["gasto"]:,.2f}'),
            ('Coordenadas', f'{c["lat"]}, {c["lng"]}'),
        ]
    )
    doc.add_paragraph('')

    add_table(doc,
        ['Parámetro', 'Valor solicitud'],
        [
            ('Monto solicitado', f'S/ {c["monto"]:,.2f}'),
            ('Plazo', f'{c["plazo"]} meses'),
            ('TEA', c['tea']),
            ('Garantía', c['garantia']),
            ('Destino', c['destino']),
            ('Cuota referencia', f'S/ {c["cuota_ref"]:,.2f}'),
            ('Prioridad cartera', c['prioridad']),
        ]
    )
    doc.add_paragraph('')

    add_table(doc,
        ['Evaluación', 'Resultado'],
        [
            ('Pre-evaluación', c['preeval']),
            ('Buró', c['buro']),
            ('Decisión comité', c['decision']),
        ]
    )
    doc.add_paragraph('')

    desemb = [
        ('Monto aprobado', f'S/ {c["monto_aprobado"]:,.2f}' if c["monto_aprobado"] else c["monto_aprobado"]),
        ('Fecha desembolso', c['fecha_desembolso']),
        ('Día pago', str(c['dia_pago'])),
        ('Cuota mensual', str(c['cuota_final']) if isinstance(c['cuota_final'], str) else f'S/ {c["cuota_final"]:,.2f}'),
    ]
    if c['decision'] == 'RECHAZADO':
        desemb = [('Resultado', 'RECHAZADO — No genera desembolso ni cronograma')]
    add_table(doc, ['Desembolso', 'Valor'], desemb)
    doc.add_paragraph('')
    doc.add_paragraph('─' * 60)
    doc.add_paragraph('')

# ═══════════════════════════ 8. RESUMEN ═══════════════════════════
doc.add_heading('8. Resumen General', level=1)
add_table(doc,
    ['Tipo', 'Cantidad'],
    [
        ('Desembolsados (aprobados)', '24'),
        ('Condicionados (monto reducido)', '3 — Casos 25, 26, 27'),
        ('Rechazados', '3 — Casos 28, 29, 30'),
    ]
)

doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph('Documento generado automáticamente el ')
p.add_run('02/07/2026').bold = True

doc.save(r'D:\backend\Credenciales_Prymera_Detallado.docx')
print('Documento creado: D:\\backend\\Credenciales_Prymera_Detallado.docx')
