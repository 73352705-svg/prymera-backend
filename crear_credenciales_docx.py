"""Generate Word document with all credentials for the Prymera ecosystem."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
import psycopg2

doc = Document()

# Styling
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Prymera — Credenciales del Sistema', level=0)
title.alignment = 1  # center

doc.add_paragraph('')

# ── PostgreSQL ──
doc.add_heading('1. Base de Datos PostgreSQL', level=1)
pg = doc.add_table(rows=5, cols=2)
pg.style = 'Light Shading Accent 1'
pg.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('Servidor', 'localhost:5432'),
    ('Base de datos', 'bd_core_mobile'),
    ('Usuario', 'postgres'),
    ('Contraseña', 'postgres'),
    ('pgAdmin', 'Registrar servidor con los datos de arriba'),
]
for i, (k, v) in enumerate(data):
    pg.rows[i].cells[0].text = k
    pg.rows[i].cells[1].text = v

doc.add_paragraph('')

# ── Backend API ──
doc.add_heading('2. Backend API REST', level=1)
api = doc.add_table(rows=4, cols=2)
api.style = 'Light Shading Accent 1'
api.alignment = WD_TABLE_ALIGNMENT.CENTER
api_data = [
    ('URL Base', 'http://localhost:8003'),
    ('Documentación', 'http://localhost:8003/docs'),
    ('Iniciar', 'D:\\backend\\iniciar_api.bat (doble click)'),
    ('Detener', 'Ctrl + C en la terminal'),
]
for i, (k, v) in enumerate(api_data):
    api.rows[i].cells[0].text = k
    api.rows[i].cells[1].text = v

doc.add_paragraph('')

# ── Asesores (Fuerza de Ventas) ──
doc.add_heading('3. Fuerza de Ventas — Asesores', level=1)

conn = psycopg2.connect(host='localhost', port=5432, dbname='bd_core_mobile', user='postgres', password='postgres')
cur = conn.cursor()
cur.execute("SELECT codigo_empleado, nombres, apellidos, perfil FROM asesores ORDER BY codigo_empleado")
asesores = cur.fetchall()
cur.close()
conn.close()

a_tbl = doc.add_table(rows=1 + len(asesores), cols=4)
a_tbl.style = 'Light Shading Accent 1'
a_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = a_tbl.rows[0].cells
hdr[0].text = 'Código'
hdr[1].text = 'Nombre'
hdr[2].text = 'Perfil'
hdr[3].text = 'Contraseña'
for i, (cod, nom, ape, perf) in enumerate(asesores):
    r = a_tbl.rows[i + 1].cells
    r[0].text = cod
    r[1].text = f'{nom} {ape}'
    r[2].text = perf
    r[3].text = 'Prymera2024'

doc.add_paragraph('')

# ── Clientes App ──
doc.add_heading('4. App Cliente — Usuarios', level=1)
doc.add_paragraph('Todos los usuarios tienen contraseña: Prymera2024')
doc.add_paragraph('')

conn = psycopg2.connect(host='localhost', port=5432, dbname='bd_core_mobile', user='postgres', password='postgres')
cur = conn.cursor()
cur.execute("""
    SELECT u.username, c.nombres, c.apellidos, c.numero_documento, c.nombre_negocio
    FROM usuarios_cliente u
    JOIN clientes c ON c.id = u.cliente_id
    ORDER BY u.username
""")
usuarios = cur.fetchall()
cur.close()
conn.close()

u_tbl = doc.add_table(rows=1 + len(usuarios), cols=5)
u_tbl.style = 'Light Shading Accent 1'
u_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = u_tbl.rows[0].cells
hdr[0].text = 'Usuario (DNI)'
hdr[1].text = 'Nombre'
hdr[2].text = 'Apellidos'
hdr[3].text = 'DNI'
hdr[4].text = 'Negocio'
for i, (usr, nom, ape, dni, neg) in enumerate(usuarios):
    r = u_tbl.rows[i + 1].cells
    r[0].text = usr
    r[1].text = nom
    r[2].text = ape
    r[3].text = dni
    r[4].text = neg or '—'

doc.add_paragraph('')

# ── Tabla resumen 30 casos ──
doc.add_heading('5. Resumen — 30 Casos de Crédito', level=1)
doc.add_paragraph(
    'Los 30 casos del archivo ENUNCIADOS_30_CASOS_CREDITO_FLUJO_MOVIL.md '
    'están disponibles en la base de datos. Para cada caso:\n'
    '• Cliente: inicia sesión en App Cliente con su DNI / Prymera2024\n'
    '• Asesor: inicia sesión en Fuerza de Ventas con 0001 / Prymera2024\n'
    '• Sigue el flujo: Solicitud → Cartera → Visita → Pre-evaluación → Buró → Documentos → Comité → Desembolso'
)

doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph('Documento generado el ')
p.add_run('2026-07-02').bold = True

doc.save(r'D:\backend\Credenciales_Prymera.docx')
print('Documento creado: D:\\backend\\Credenciales_Prymera.docx')
